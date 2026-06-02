"""
Service layer for confluence integration in the CortexSail Agentic RAG System
"""

from src.configuration import settings
from src.crew.crewmanager import CrewManager
from typing import Dict, List
from src.memory.conversation_memory import get_conversation_memory
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import WD_ALIGN_VERTICAL

class ConfluenceService:

    """Service layer for handling interactions with Confluence API"""

    def __init__(self):
        ##crew_manager - Workspace brain holds the context and state for the current workspace, including Confluence credentials and session management - its for a single conversation or session
        self.crew_manager = None
        self.current_session_id = None
        self.conversation_memory = None
    
    def user_prompt(self, query: str) -> Dict:
        """Analyse user query and determine if it requires Confluence interaction, then route to appropriate agent or tool"""

        try:
            print(f"Extract the confluence credentials from the session")

            ###Fix this on actual session management implementation with auth services and session handling
            session_id = "default session"

            if self.crew_manager is None or self.current_session_id != session_id:

            ##if session not there and crew manager not initialized, initialize its object and the condition to check if the session has changed, if so, reinitialize the crew manager for the new session
                print (f"Initialize the cew manager for the session {session_id}")

                self.crew_manager = CrewManager(session_id=session_id)

                self.current_session_id = session_id

            ###Initializing the conversaion memory for the session - this is temporary until we have a proper session management and memory management in place

            if self.conversation_memory is None:
                print(f"Creating conversation memory for the session {session_id}")

                self.conversation_memory = get_conversation_memory(session_id)

            ###Fetching the conversation summary for the session from the conversation memory, this can be used to provide context to the agents for the conversation - latest conversation entry

            last_conversation = self.conversation_memory.get_latest_data_from_latest_query()

            #Fetching filters also
            conversation_filter = self.conversation_memory.get_latest_filters()

            print(f"Fetched previous data: {len(last_conversation) if last_conversation else 0}")

            ###Initializing the crew instance
            crew_instance = self.crew_manager.get_crew_instance()

            ##Collecting the context and routing to crew instance , LLM will determine the next steps and the tools to use based on the query, conversation history and filters 

            crew_inputs_llm = {
                "prompt": query,
                "previous_conversation":
                last_conversation if last_conversation else [],
                "previous_filters": conversation_filter if conversation_filter else {}
            }

            result = crew_instance.kickoff(input = crew_inputs_llm)

            ##Extracting the response
            if hasattr(result, "raw"):
                output_text = result.raw
            elif isinstance(result, dict) and "raw" in result:
                output_text = result["raw"]
            else:
                output_text = str(result)

            ##If no data is returned, 
            no_data_response = [
                "No relevant documents found",
                "No relevant information found",
                "No relevant data found",
                "No relevant results found",
                "No relevant content found",
                "No data found",
                "No information found",
                "No results found",
                "No content found",
                "could not find any relevant information",
                "could not find any relevant documents",
                "could not find any relevant data",
            ]

            if output_text.strip() == "" or any(phrase in output_text.lower() for phrase in no_data_response):
                return {
                    "success": False,
                    "output": "",
                    "message": "No relevant documents found matching your criteria"
                }
            print(f"Document retrieval successful, returning response to user")

            ##Storing this reponse in the conversation memory for the session, this can be used to provide context to the agents for the conversation
            try:
                result_data = {
                    "intent": getattr(result, "intent","analysis"),
                    "filters": crew_inputs_llm.get("previous_filters",{}),
                    "confluence_document": getattr(result, "confluence_document", []),
                    "output_format": getattr(result, "output_format", {"text": output_text}),
                }   
                ##Fallback mechanism
                ##task_output extraction from result, fallback if the result does not contain these fields then we use the crew generated task_output list - from kickoff

                if hasattr(result, "tasks_output") and result.tasks_output:
                    for task_output in result.tasks_output:
                        ##Parsing JSON from task output
                        import json
                        try:
                            parsed_data = json.loads(task_output.raw)
                            if "intent" in parsed_data:
                                result_data["intent"] = parsed_data["intent"]
                            if "filters" in parsed_data:
                                result_data["filters"] = parsed_data["filters"]
                            if "confluence_document" in parsed_data:
                                result_data["confluence_document"] = parsed_data["confluence_document"]
                            if "output_format" in parsed_data:
                                result_data["output_format"] = parsed_data["output_format"]
                        except:
                            pass
                self.conversation_memory.add_to_conversation_history(query,result_data)
                print(f"Stored the response in the conversation memory for the session")
            except Exception as memory_err:
                print(f"Unable to store the response in the conversation memory: {memory_err}")

            ##Return object
            return {
                "success": True,
                "output": output_text,
                "message": "Memory process completed successfully"
            }
        except Exception as e:
            print(f"Error occured during the process: {e}")
            return {
                "success": False,
                "output": "",
                "message": f"An error occurred: {str(e)}"
            }
    
    ###Function definintion for generating the word document from the markdown 
    def produce_word_document_from_markdown(self, markdown_content: str, session_id: str) -> Dict:
        """Generate a word document from markdown content and return the word file or the path URL
        
        Args:
            markdown_content (str): The markdown content to be converted to a word document
            session_id (str): The session ID for which the document is being generated, this can be used for storing the document in a session specific location or for associating the document with the session in the memory

        Returns:
            Dict: A dictionary containing the success status, the output (word document or URL), and a message
        """

        try:
            #Create a document
            doc = Document()  
            
            ##Title
            title = doc.add_heading("Confluence Document Analysis",0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            ##Timestamp
            timestamp = doc.add_paragraph(
                f"Generated at: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}"
            )

            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            ##Content Splitting for processing
            lines = markdown_content.split("\n")

            i=0;
            while i < len(lines):
                line = lines[i].strip()

                if not line:
                    i+=1
                    continue
                
                ##To replace the #'s with the appropriate heading levels by replacing it
                elif line.startswith("###"):
                    doc.add_heading(line.replace("###","").strip(), level=1)
                elif line.startswith("####"):
                    doc.add_heading(line.replace("####","").strip(), level=2)
                elif line.startswith("#####"):
                    doc.add_heading(line.replace("#####","").strip(), level=3)
                
                ###Handling tables
                elif line.startswith("|"):
                    table_lines = []
                    while i<len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i].strip())
                        i+=1
                ##Skipping the header separator line in markdown tables

                ## Remove the separator row that is present in markdown tables by default after the header row
                if table_lines and len(table_lines)>1 and '---' in table_lines[1]:
                    table_lines.pop(1)
                
                #Creating table
                if table_lines:
                    self.create_table_markdown_lines(doc, table_lines)
                continue

            ##Handling bullet points
            elif line.startswith("- "):
                bullet_point = line.replace("- ","",1)
                ##Handling bold text and storing in the bullet point
                bullet_point = self._sanitize_bold_markdown(bullet_point)
                para = doc.add_paragraph(bullet_point, style='List Bullet')
            
            ##Handling the horizontal lines in the markdown and adding a horizontal line in the word document and centering it
            elif line.startswith("---"):
                doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

            ##blockquotes in markdown
            elif line.startswith("> "):
                blockquote_text = line.replace(">  ", "", 1)
                para = doc.add_paragraph(blockquote_text)
                para.italic = True

            

            



    ##Used inside the produce_word_document_from_markdown
    def create_table_markdown_lines(self, doc, table_lines):
        """Word document generator for adding tables from markdown lines

        Args:
            doc (Document): The word document object to which the table will be added
            table_lines (List[str]): The lines of markdown content that represent the table, these lines will be parsed to create the table in the word document
            
        Returns:
            None: This function modifies the word document object in place and does not return anything
        """
        rows = []
        if not table_lines:
            return
        ##Parsing the table data lines
        for line in table_lines:
            ##Here removing the first and last character which are | in markdown tables and then splitting the line by | to get the individual cells also replacing the space after the split and removing the 1th index and the -1th index
            cells_list = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells_list:
                rows.append(cells_list)

            if not rows:
                return
            ##Creating the table with the rows[]
            ##Choosing the number of elements of any rows to get the count of the columns
            #fixed rows and cols
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'

            ##Filling the data - update the n^2 complexity here and optimize it

            #Enumerate gives the index and the entire first row here 1st iteration 
            for row_index, row_data in enumerate(rows):
                for col_index, cell_data in enumerate(row_data):
                    cell = table.cell(row_index, col_index)
                    cell.text = self._clean_text(cell_data)

                    #Just header row needs to be bold
                    ###cell -> paragraph -> run -> (Text that requires formatting like bold, italic)
                    if row_index == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    ##vertical alignment for the cells 
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ##bold markdown fix
    def _clean_text(self, text):
        """Helper function to clean the text by removing the bold markers ** """
        text = text.replace("**","")
        return text
    ##Helper function for bold text
    def _sanitize_bold_markdown(self, text):
        """Helper function to sanitize the bold markdown in the text by removing the ** markers and returning the cleaned text
        """
        return text.replace("**","")

